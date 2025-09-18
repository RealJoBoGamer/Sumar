import os
import json
import tkinter as tk
from tkinter import filedialog, messagebox, scrolledtext, simpledialog
from docx import Document
from openai import OpenAI
import requests
import msal

# ====== Config & Cache ======
CONFIG_FILE = "config.json"
CACHE_FILE = "token_cache.json"

def load_config():
    if os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, "r") as f:
            return json.load(f)
    return {}

def save_config(cfg):
    with open(CONFIG_FILE, "w") as f:
        json.dump(cfg, f, indent=2)

def load_cache():
    cache = msal.SerializableTokenCache()
    if os.path.exists(CACHE_FILE):
        with open(CACHE_FILE, "r") as f:
            cache.deserialize(f.read())
    return cache

def save_cache(cache):
    if cache.has_state_changed:
        with open(CACHE_FILE, "w") as f:
            f.write(cache.serialize())

# ====== Microsoft Graph setup ======
CLIENT_ID = "0a8aff2a-148b-4b77-9d18-af40f118ccce"   # Azure App Registration
TENANT_ID = "c356ada5-b3f7-4a81-bee7-1657c780ad12"
AUTHORITY = f"https://login.microsoftonline.com/{TENANT_ID}"
SCOPES = ["User.Read", "Files.ReadWrite"]

def get_access_token(root):
    """Authenticate with Microsoft Graph using device code flow"""
    cache = load_cache()
    app = msal.PublicClientApplication(CLIENT_ID, authority=AUTHORITY, token_cache=cache)
    accounts = app.get_accounts()

    if accounts:
        result = app.acquire_token_silent(SCOPES, account=accounts[0])
    else:
        flow = app.initiate_device_flow(scopes=SCOPES)
        if "user_code" not in flow:
            raise Exception("Device flow initiation failed")
        messagebox.showinfo("Microsoft Login", flow["message"])
        result = app.acquire_token_by_device_flow(flow)

    save_cache(cache)

    if "access_token" in result:
        return result["access_token"]
    else:
        raise Exception(f"Authentication failed: {result}")

def list_drive_items(token, folder_id=None):
    """List items in OneDrive folder"""
    headers = {"Authorization": f"Bearer {token}"}
    if folder_id:
        url = f"https://graph.microsoft.com/v1.0/me/drive/items/{folder_id}/children"
    else:
        url = "https://graph.microsoft.com/v1.0/me/drive/root/children"
    resp = requests.get(url, headers=headers).json()
    return resp.get("value", [])

def download_file(token, file_id, save_as):
    """Download OneDrive file by ID"""
    headers = {"Authorization": f"Bearer {token}"}
    url = f"https://graph.microsoft.com/v1.0/me/drive/items/{file_id}/content"
    resp = requests.get(url, headers=headers, stream=True)
    with open(save_as, "wb") as f:
        for chunk in resp.iter_content(chunk_size=8192):
            f.write(chunk)

def upload_file(token, file_path):
    """Upload file back to OneDrive root"""
    headers = {"Authorization": f"Bearer {token}"}
    filename = os.path.basename(file_path)
    url = f"https://graph.microsoft.com/v1.0/me/drive/root:/{filename}:/content"
    with open(file_path, "rb") as f:
        resp = requests.put(url, headers=headers, data=f)
    return resp.json()

# ====== Summarization logic ======
def read_word_documents(directory):
    """Read .docx files from a local folder"""
    text_content = []
    for filename in os.listdir(directory):
        if filename.endswith(".docx"):
            filepath = os.path.join(directory, filename)
            doc = Document(filepath)
            for para in doc.paragraphs:
                if para.text.strip():
                    text_content.append(para.text)
    return "\n".join(text_content)

def summarise_text(client, text):
    """Summarize via OpenRouter"""
    response = client.chat.completions.create(
        model="mistralai/mistral-7b-instruct:free",
        messages=[
            {"role": "system",
             "content": (
                 "You are a helpful assistant that summarizes documents into concise note form. "
                 "If there are errors, missing details, or lessons to be learned, identify them and fill in the gaps."
             )},
            {"role": "user", "content": text}
        ]
    )
    return response.choices[0].message.content.strip()

# ====== GUI Application ======
class SummarizerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("Sumar (Self-Contained)")
        self.root.geometry("800x600")

        # Config
        self.config = load_config()
        if "openrouter_api_key" not in self.config:
            key = simpledialog.askstring("OpenRouter Key", "Enter your OpenRouter API Key:")
            if key:
                self.config["openrouter_api_key"] = key
                save_config(self.config)
            else:
                messagebox.showerror("Error", "API key required.")
                root.destroy()
                return
        self.client = OpenAI(
            base_url="https://openrouter.ai/api/v1",
            api_key=self.config["openrouter_api_key"]
        )

        # Buttons
        self.btn_select_local = tk.Button(root, text="Use Local Folder", command=self.select_local)
        self.btn_select_local.pack(pady=5)

        self.btn_select_onedrive = tk.Button(root, text="Use OneDrive Folder", command=self.select_onedrive)
        self.btn_select_onedrive.pack(pady=5)

        self.btn_change_key = tk.Button(root, text="Change OpenRouter Key", command=self.change_key)
        self.btn_change_key.pack(pady=5)

        self.btn_signout = tk.Button(root, text="Sign Out of Microsoft", command=self.sign_out)
        self.btn_signout.pack(pady=5)

        self.btn_select_save = tk.Button(root, text="Select Save Location", command=self.select_save)
        self.btn_select_save.pack(pady=5)

        self.btn_run = tk.Button(root, text="Run Summarizer", command=self.run_summarizer, state=tk.DISABLED)
        self.btn_run.pack(pady=5)

        # Output log
        self.output = scrolledtext.ScrolledText(root, wrap=tk.WORD, height=20, width=100, state=tk.DISABLED)
        self.output.pack(pady=10)

        # Paths
        self.directory = None
        self.temp_downloads = []
        self.save_path = None
        self.token = None

    def log(self, message):
        self.output.config(state=tk.NORMAL)
        self.output.insert(tk.END, message + "\n")
        self.output.see(tk.END)
        self.output.config(state=tk.DISABLED)
        self.root.update()

    def select_local(self):
        self.directory = filedialog.askdirectory()
        if self.directory:
            self.log(f"Selected local folder: {self.directory}")
        self.check_ready()

    def select_onedrive(self):
        try:
            self.token = get_access_token(self.root)
            folder_id = None
            while True:
                items = list_drive_items(self.token, folder_id)
                folders = [i for i in items if i["folder"]["childCount"] > 0] if items else []
                files = [i for i in items if i["name"].endswith(".docx")] if items else []

                choice = simpledialog.askstring(
                    "OneDrive Browser",
                    "Folders:\n" + "\n".join([f"[{idx}] {f['name']}" for idx, f in enumerate(folders)]) +
                    "\n\nFiles:\n" + "\n".join([f"- {f['name']}" for f in files]) +
                    "\n\nEnter folder index to navigate, or leave blank to select current folder."
                )

                if choice and choice.isdigit() and int(choice) < len(folders):
                    folder_id = folders[int(choice)]["id"]
                else:
                    break

            self.log("Downloading OneDrive docs...")
            os.makedirs("onedrive_temp", exist_ok=True)
            self.temp_downloads = []
            for file in files:
                local_path = os.path.join("onedrive_temp", file["name"])
                download_file(self.token, file["id"], local_path)
                self.temp_downloads.append(local_path)
            self.log("OneDrive docs downloaded successfully.")
        except Exception as e:
            messagebox.showerror("OneDrive Error", str(e))
        self.check_ready()

    def select_save(self):
        self.save_path = filedialog.asksaveasfilename(
            defaultextension=".txt",
            filetypes=[("Text Files", "*.txt")],
            title="Save summary as"
        )
        if self.save_path:
            self.log(f"Save path: {self.save_path}")
        self.check_ready()

    def check_ready(self):
        if (self.directory or self.temp_downloads) and self.save_path:
            self.btn_run.config(state=tk.NORMAL)

    def run_summarizer(self):
        if self.directory:
            self.log("Reading local Word documents...")
            text = read_word_documents(self.directory)
        elif self.temp_downloads:
            self.log("Reading downloaded OneDrive documents...")
            text_content = []
            for filepath in self.temp_downloads:
                doc = Document(filepath)
                for para in doc.paragraphs:
                    if para.text.strip():
                        text_content.append(para.text)
            text = "\n".join(text_content)
        else:
            messagebox.showerror("Error", "No documents selected.")
            return

        if not text.strip():
            messagebox.showerror("Error", "No text found in documents.")
            return

        self.log("Sending text to AI summarizer... please wait.")
        try:
            summary = summarise_text(self.client, text)
        except Exception as e:
            messagebox.showerror("Error", f"AI summarization failed:\n{e}")
            return

        self.log("Writing summary to file...")
        with open(self.save_path, "w", encoding="utf-8") as f:
            f.write(summary)

        self.log("✅ Done! Summary saved successfully.")
        messagebox.showinfo("Success", f"Summary saved at:\n{self.save_path}")

        if self.token and self.temp_downloads:
            if messagebox.askyesno("Upload", "Upload summary back to OneDrive?"):
                try:
                    result = upload_file(self.token, self.save_path)
                    self.log(f"Uploaded summary to OneDrive: {result.get('name', 'Unknown')}")
                except Exception as e:
                    messagebox.showerror("Upload Error", str(e))

    def sign_out(self):
        if os.path.exists(CACHE_FILE):
            os.remove(CACHE_FILE)
        messagebox.showinfo("Sign Out", "Signed out successfully.")

    def change_key(self):
        key = simpledialog.askstring("Change Key", "Enter new OpenRouter API Key:")
        if key:
            self.config["openrouter_api_key"] = key
            save_config(self.config)
            self.client = OpenAI(
                base_url="https://openrouter.ai/api/v1",
                api_key=key
            )
            messagebox.showinfo("Updated", "API Key updated successfully.")

# ====== Main ======
def main():
    root = tk.Tk()
    app = SummarizerApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()
